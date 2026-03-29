#!/usr/bin/env python3
"""
doc2graph.py — Estrae un grafo di conoscenza da qualsiasi documento
e genera un visualizzatore HTML/CSS/JS interattivo stile Neo4j.

Formati supportati: .txt .md .pdf .docx .doc .epub .odt .rtf .csv .json
Utilizza un LLM OpenAI-compatible (LM-Studio, Ollama, OpenAI…) via .env

MULTI-FILE: passa più file o glob come argomenti → unico grafo unificato.
CHECKPOINT: salva il progresso dopo ogni file, riprende se interrotto.
MERGE MODE: --merge-jsons a.json b.json -o finale.html  (zero LLM)
"""

import os
import sys
import json
import re
import gc
import argparse
import subprocess
import tempfile
import unicodedata
from pathlib import Path
from datetime import datetime

# ── third-party ─────────────────────────────────────────────────────────────
try:
    from dotenv import load_dotenv
except ImportError:
    print("❌  Manca python-dotenv: pip install python-dotenv")
    sys.exit(1)

try:
    from openai import OpenAI
except ImportError:
    print("❌  Manca openai: pip install openai")
    sys.exit(1)

# ─────────────────────────────────────────────────────────────────────────────
# 1. CONFIG  (.env)
# ─────────────────────────────────────────────────────────────────────────────

load_dotenv()

LLM_BASE_URL   = os.getenv("LLM_BASE_URL",   "http://localhost:1234/v1")
LLM_API_KEY    = os.getenv("LLM_API_KEY",    "lm-studio")
LLM_MODEL      = os.getenv("LLM_MODEL",      "").strip()   # auto-discover se vuoto
LLM_MAX_TOKENS = int(os.getenv("LLM_MAX_TOKENS", "4096"))
CHUNK_SIZE     = int(os.getenv("CHUNK_SIZE",  "6000"))
CHUNK_OVERLAP  = int(os.getenv("CHUNK_OVERLAP", "500"))
LLM_RETRY      = int(os.getenv("LLM_RETRY",  "3"))        # tentativi su risposta vuota

client = OpenAI(base_url=LLM_BASE_URL, api_key=LLM_API_KEY)


def _resolve_model() -> str:
    global LLM_MODEL
    if LLM_MODEL:
        return LLM_MODEL

    print("⚠️  LLM_MODEL non impostato nel .env — auto-discovery modelli…")
    try:
        models = client.models.list()
        available = [m.id for m in models.data]
        if not available:
            print("❌  Nessun modello disponibile sul server LLM.")
            print("   Imposta LLM_MODEL nel .env e riprova.")
            sys.exit(1)
        LLM_MODEL = available[0]
        print(f"   ✅  Modelli disponibili: {available}")
        print(f"   ✅  Uso automaticamente: {LLM_MODEL}")
        return LLM_MODEL
    except Exception as e:
        print(f"❌  Impossibile interrogare {LLM_BASE_URL}/models: {e}")
        print("   Controlla che LM-Studio (o il tuo server LLM) sia avviato e raggiungibile.")
        sys.exit(1)

# ─────────────────────────────────────────────────────────────────────────────
# 2. ESTRAZIONE TESTO  (multi-formato)
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(path: Path) -> str:
    ext = path.suffix.lower()

    if ext in {".txt", ".md", ".log", ".rst", ".yaml", ".yml", ".toml",
               ".py", ".js", ".ts", ".java", ".cs", ".cpp", ".c", ".go"}:
        return path.read_text(encoding="utf-8", errors="replace")

    if ext == ".pdf":
        return _extract_pdf(path)

    if ext == ".docx":
        return _extract_docx(path)

    if ext in {".csv", ".tsv"}:
        return _extract_csv(path, ext)

    if ext in {".json", ".jsonl"}:
        return _extract_json(path, ext)

    if ext in {".epub", ".odt", ".rtf", ".doc", ".pptx"}:
        return _pandoc_to_text(path)

    print(f"⚠️  Formato '{ext}' non nativamente supportato — tentativo con pandoc…")
    return _pandoc_to_text(path)


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append(f"--- Pagina {i+1} ---\n{text}")
        full = "\n\n".join(pages)
        if len(full.strip()) < 200:
            result = subprocess.run(
                ["pdftotext", "-layout", str(path), "-"],
                capture_output=True, text=True
            )
            if result.returncode == 0:
                full = result.stdout
        return full
    except Exception as e:
        return f"[Errore lettura PDF: {e}]"


def _extract_docx(path: Path) -> str:
    try:
        import docx as _docx
        doc = _docx.Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name if para.style else ""
                prefix = ""
                if "Heading" in style:
                    level = style.replace("Heading ", "")
                    prefix = "#" * int(level) + " " if level.isdigit() else "## "
                parts.append(prefix + para.text)
        for idx, table in enumerate(doc.tables):
            parts.append(f"\n[TABELLA {idx+1}]")
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                header = rows[0]
                sep = "| " + " | ".join(["---"] * len(table.columns)) + " |"
                parts.append(header)
                parts.append(sep)
                parts.extend(rows[1:])
            parts.append("")
        return "\n".join(parts)
    except Exception:
        return _pandoc_to_text(path)


def _extract_csv(path: Path, ext: str) -> str:
    try:
        import csv
        sep = "\t" if ext == ".tsv" else ","
        with open(path, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f, delimiter=sep)
            rows = list(reader)
        if not rows:
            return ""
        header = rows[0]
        lines = ["| " + " | ".join(header) + " |",
                 "| " + " | ".join(["---"] * len(header)) + " |"]
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return f"[CSV — {len(rows)-1} righe, {len(header)} colonne]\n" + "\n".join(lines)
    except Exception as e:
        return f"[Errore lettura CSV: {e}]"


def _extract_json(path: Path, ext: str) -> str:
    try:
        if ext == ".jsonl":
            lines = []
            with open(path, encoding="utf-8", errors="replace") as f:
                for line in f:
                    try:
                        lines.append(json.loads(line))
                    except Exception:
                        pass
            return json.dumps(lines[:50], indent=2, ensure_ascii=False)
        else:
            data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
            return json.dumps(data, indent=2, ensure_ascii=False)
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")


def _pandoc_to_text(path: Path) -> str:
    try:
        result = subprocess.run(
            ["pandoc", str(path), "-t", "plain", "--wrap=none"],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            return result.stdout
        return f"[Pandoc error: {result.stderr[:300]}]"
    except FileNotFoundError:
        return f"[pandoc non trovato — installa pandoc per supportare {path.suffix}]"
    except Exception as e:
        return f"[Errore pandoc: {e}]"


# ─────────────────────────────────────────────────────────────────────────────
# 3. CHUNKING con overlap
# ─────────────────────────────────────────────────────────────────────────────

def split_into_chunks(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        if end < len(text):
            cut = text.rfind("\n\n", start, end)
            if cut == -1:
                cut = text.rfind("\n", start, end)
            if cut != -1 and cut > start:
                end = cut
        chunks.append(text[start:end].strip())
        if end >= len(text):
            break
        start = end - overlap
    return [c for c in chunks if c]


# ─────────────────────────────────────────────────────────────────────────────
# 4. PROMPT DI ESTRAZIONE
# ─────────────────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """Sei un esperto di knowledge graph e Neo4j.
Il tuo compito e' analizzare testo ed estrarne un grafo di conoscenza strutturato.

Rispondi ESCLUSIVAMENTE con un JSON valido, senza markdown, senza backtick, senza spiegazioni.

Struttura obbligatoria:
{
  "nodes": [
    {
      "id": "n1",
      "label": "NomeNodo",
      "type": "Categoria",
      "properties": { "chiave": "valore" },
      "description": "Breve descrizione opzionale"
    }
  ],
  "edges": [
    {
      "source": "n1",
      "target": "n2",
      "type": "TIPO_RELAZIONE",
      "label": "etichetta leggibile breve",
      "properties": { "dal": "2020", "ruolo": "CEO" },
      "evidence": "frase testuale che giustifica questa relazione"
    }
  ]
}

REGOLE SUI NODI:
- Tipi concisi: Persona, Organizzazione, Concetto, Prodotto, Evento, Luogo, Tecnologia, Legge, Dato, Documento
- Estrai nelle properties SOLO valori esplicitamente citati (date, numeri, ruoli, stati)
- Max 5 properties per nodo
- Gli id devono essere stringhe univoche brevi (n1, n2, ...)

REGOLE CRITICHE SULLE RELAZIONI:

1. MAI usare relazioni generiche. Sono VIETATE: RELAZIONATO_A, CONNESSO_A, ASSOCIATO_A,
   COLLEGATO_A, MENZIONA, HA_RELAZIONE, E_CONNESSO, APPARTIENE.
   Ogni relazione deve descrivere il VERBO PRECISO del testo.

2. Il campo "type" deve essere il verbo in MAIUSCOLO_UNDERSCORE ricavato letteralmente dal testo.
   Esempi CORRETTI:
   - "Mario ha fondato Acme" -> FONDA
   - "La legge vieta l'uso di X" -> VIETA_USO_DI
   - "Il sistema elabora i dati" -> ELABORA
   - "L'azienda acquisisce il competitor" -> ACQUISISCE
   - "Il contratto scade nel 2025" -> SCADE_IN
   - "Il brevetto e' stato depositato da Rossi" -> DEPOSITA
   Esempi SBAGLIATI (non usare mai): RELAZIONATO_A, HA_RELAZIONE, E_CONNESSO, APPARTIENE

3. Il campo "label" e' una frase brevissima (max 5 parole) leggibile da umano, es:
   - "ha fondato nel 1998"
   - "vieta l'utilizzo commerciale"
   - "acquisisce per 2 miliardi"
   - "elabora in tempo reale"

4. Nelle "properties" degli edge metti i dettagli quantitativi della relazione se presenti:
   data, durata, importo, modalita', condizione, percentuale, ruolo specifico.

5. Includi SOLO relazioni ESPLICITAMENTE dichiarate nel testo.
   Ogni edge DEVE avere "evidence": citazione testuale diretta che la giustifica.
   Se non trovi la frase -> NON inserire l'edge.

6. NON aggiungere IS_A, PART_OF, INSTANCE_OF se non scritti nel testo.
"""

ENRICH_PROMPT = """Sei un esperto di knowledge graph.
Ricevi un JSON con nodes e edges estratti da un documento.

Alcuni edges hanno type generico (RELAZIONATO_A, CONNESSO_A, ASSOCIATO_A, HA_RELAZIONE ecc.)
oppure label vuota o poco descrittiva.

Il tuo compito:
1. Per ogni edge con type generico: sostituiscilo con il verbo preciso ricavabile dall'evidence.
2. Per ogni edge con label vuota o generica: scrivi una label leggibile di max 5 parole.
3. Per ogni edge: aggiungi properties se l'evidence contiene dettagli (date, importi, ruoli, condizioni).
4. NON aggiungere, rimuovere o spostare nodi o edges: migliora solo type, label, properties.

Tipi generici DA SOSTITUIRE: RELAZIONATO_A, CONNESSO_A, ASSOCIATO_A, COLLEGATO_A,
MENZIONA, HA_RELAZIONE, E_CONNESSO, APPARTIENE, RIFERITO_A, LEGATO_A.

Rispondi ESCLUSIVAMENTE con il JSON aggiornato, stesso formato dell'input, senza markdown.
"""

# ─────────────────────────────────────────────────────────────────────────────
# 5. PARSING JSON ROBUSTO
# ─────────────────────────────────────────────────────────────────────────────

def repair_json(raw: str) -> str:
    s = raw.strip()
    s = re.sub(r"^```[a-zA-Z]*\s*", "", s)
    s = re.sub(r"\s*```\s*$", "", s)
    s = s.strip()
    m = re.search(r"(\{[\s\S]*\}|\[[\s\S]*\])", s)
    if m:
        s = m.group(1)
    s = _remove_comments(s)
    s = re.sub(r",(\s*[}\]])", r"\1", s)
    s = re.sub(r'\bNone\b',  "null",  s)
    s = re.sub(r'\bTrue\b',  "true",  s)
    s = re.sub(r'\bFalse\b', "false", s)
    if '"' not in s:
        s = s.replace("'", '"')
    s = re.sub(r'([{,]\s*)([a-zA-Z_][a-zA-Z0-9_]*)\s*:', r'\1"\2":', s)
    s = _close_truncated(s)
    return s


def _remove_comments(s: str) -> str:
    result = []
    i = 0
    in_string = False
    while i < len(s):
        c = s[i]
        if c == '\\' and in_string:
            result.append(c)
            result.append(s[i+1] if i+1 < len(s) else '')
            i += 2
            continue
        if c == '"':
            in_string = not in_string
            result.append(c)
            i += 1
            continue
        if not in_string:
            if s[i:i+2] == '//':
                while i < len(s) and s[i] != '\n':
                    i += 1
                continue
            if s[i:i+2] == '/*':
                end = s.find('*/', i+2)
                i = end + 2 if end != -1 else len(s)
                continue
        result.append(c)
        i += 1
    return ''.join(result)


def _close_truncated(s: str) -> str:
    stack = []
    in_string = False
    escape_next = False
    for ch in s:
        if escape_next:
            escape_next = False
            continue
        if ch == '\\' and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            stack.append('}')
        elif ch == '[':
            stack.append(']')
        elif ch == '}' and stack and stack[-1] == '}':
            stack.pop()
        elif ch == ']' and stack and stack[-1] == ']':
            stack.pop()
    if stack:
        s = re.sub(r",\s*$", "", s.rstrip())
        s += ''.join(reversed(stack))
    return s


def _extract_partial(raw: str) -> dict:
    nodes, edges = [], []
    nid = 0
    q = r'["\']'
    label_re  = re.compile(q + r'label'  + q + r'\s*:\s*' + q + r'([^"\']+)' + q)
    type_re   = re.compile(q + r'type'   + q + r'\s*:\s*' + q + r'([^"\']+)' + q)
    source_re = re.compile(q + r'source' + q + r'\s*:\s*' + q + r'([^"\']+)' + q)
    target_re = re.compile(q + r'target' + q + r'\s*:\s*' + q + r'([^"\']+)' + q)
    blocks = re.split(r'\},?\s*\{', raw)
    for block in blocks:
        lm  = label_re.search(block)
        tm  = type_re.search(block)
        sm  = source_re.search(block)
        trm = target_re.search(block)
        if sm and trm:
            edges.append({
                "source":     sm.group(1),
                "target":     trm.group(1),
                "type":       tm.group(1) if tm else "RELAZIONATO_A",
                "properties": {},
                "evidence":   "",
            })
        elif lm:
            nid += 1
            nodes.append({
                "id":          f"n{nid}",
                "label":       lm.group(1),
                "type":        tm.group(1) if tm else "Nodo",
                "properties":  {},
                "description": "",
            })
    return {"nodes": nodes, "edges": edges}


def safe_parse_llm_json(raw: str) -> dict | None:
    try:
        return json.loads(raw.strip())
    except json.JSONDecodeError:
        pass

    s1 = re.sub(r"^```[a-zA-Z]*\s*", "", raw.strip())
    s1 = re.sub(r"\s*```\s*$", "", s1).strip()
    m = re.search(r"(\{[\s\S]*\}|\[[\s\S]*\])", s1)
    if m:
        s1 = m.group(1)
    try:
        return json.loads(s1)
    except json.JSONDecodeError as e:
        print(f"\n      [strat1] {e}")

    try:
        s2 = repair_json(raw)
        return json.loads(s2)
    except json.JSONDecodeError as e:
        print(f"\n      [strat2] {e}")

    try:
        s3 = repair_json(raw)
        m2 = re.search(r"(\{[\s\S]*\})", s3)
        if m2:
            return json.loads(m2.group(1))
    except json.JSONDecodeError as e:
        print(f"\n      [strat3] {e}")

    try:
        result = _extract_partial(raw)
        if result["nodes"] or result["edges"]:
            print(f"\n      [strat4 regex] {len(result['nodes'])} nodi, {len(result['edges'])} archi")
            return result
    except Exception as e:
        print(f"\n      [strat4] {e}")

    print(f"\n      Raw (primi 400 chars): {repr(raw[:400])}")
    return None


# ─────────────────────────────────────────────────────────────────────────────
# 6. ESTRAZIONE GRAFO via LLM
# ─────────────────────────────────────────────────────────────────────────────

def _call_llm(messages: list, max_tokens: int, temperature: float = 0.1) -> str:
    import time
    model = _resolve_model()

    for attempt in range(1, LLM_RETRY + 1):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                max_tokens=max_tokens,
                temperature=temperature,
            )
            raw = response.choices[0].message.content

            if not raw or not raw.strip():
                finish = response.choices[0].finish_reason or "?"
                print(f"\n   ⚠️  Risposta vuota (tentativo {attempt}/{LLM_RETRY}) finish_reason={finish}")
                if finish == "length":
                    print(f"      → max_tokens ({max_tokens}) troppo basso. "
                          f"Aumenta LLM_MAX_TOKENS nel .env")
                elif finish in ("stop", None, "?"):
                    print(f"      → Il modello non ha generato output.")
                if attempt < LLM_RETRY:
                    wait = 2 ** attempt
                    print(f"      → Ritento tra {wait}s...", flush=True)
                    time.sleep(wait)
                continue

            return raw.strip()

        except Exception as e:
            err_str = str(e)
            print(f"\n   ❌  Errore LLM (tentativo {attempt}/{LLM_RETRY}): {err_str}")

            if "Connection refused" in err_str or "connect" in err_str.lower():
                print(f"      → Server non raggiungibile su {LLM_BASE_URL}")
            elif "model" in err_str.lower() and "not found" in err_str.lower():
                print(f"      → Modello '{model}' non trovato.")
                try:
                    avail = [m.id for m in client.models.list().data]
                    for mid in avail:
                        print(f"           • {mid}")
                except Exception:
                    pass
            elif "context" in err_str.lower() or "too long" in err_str.lower():
                print(f"      → Testo troppo lungo. Riduci CHUNK_SIZE nel .env")

            if attempt < LLM_RETRY:
                wait = 2 ** attempt
                print(f"      → Ritento tra {wait}s…", flush=True)
                time.sleep(wait)

    return ""


def llm_extract_graph(text_chunk: str, chunk_idx: int, total: int) -> dict:
    print(f"   🤖  Analisi chunk {chunk_idx}/{total} ({len(text_chunk)} chars)…",
          end=" ", flush=True)

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user",   "content": (
            "Analizza il seguente testo ed estrai il grafo di conoscenza.\n"
            "Ricorda: includi un edge SOLO se hai una \"evidence\" testuale diretta.\n\n"
            f"TESTO:\n{text_chunk}"
        )},
    ]

    raw = _call_llm(messages, LLM_MAX_TOKENS)

    if not raw:
        print("❌  Nessuna risposta dal LLM — chunk saltato")
        return {"nodes": [], "edges": []}

    data = safe_parse_llm_json(raw)
    if data is None:
        print("❌  Impossibile parsare JSON — chunk saltato")
        return {"nodes": [], "edges": []}

    nodes = data.get("nodes", [])
    edges = data.get("edges", [])
    print(f"✅  {len(nodes)} nodi, {len(edges)} archi")
    return data


# ─────────────────────────────────────────────────────────────────────────────
# 7. NORMALIZZAZIONE LABEL
# ─────────────────────────────────────────────────────────────────────────────

def normalize_label(label: str) -> str:
    s = label.strip().lower()
    s = unicodedata.normalize("NFKD", s)
    s = re.sub(r"[^\w\s]", "", s)
    s = re.sub(r"\s+", " ", s)
    return s.strip()


# ─────────────────────────────────────────────────────────────────────────────
# 8. MERGE GRAFI
# ─────────────────────────────────────────────────────────────────────────────

def merge_graphs(graph_list: list[dict]) -> dict:
    all_nodes: dict[str, dict] = {}
    all_edges: list[dict] = []
    counter = 0

    for g in graph_list:
        local_map: dict[str, str] = {}

        for node in g.get("nodes", []):
            orig_id = node.get("id", "")
            label   = (node.get("label") or "").strip()
            ntype   = (node.get("type")  or "Nodo").strip()
            key     = f"{normalize_label(label)}::{ntype.lower()}"

            if key not in all_nodes:
                counter += 1
                global_id = f"n{counter}"
                all_nodes[key] = {
                    "id":          global_id,
                    "label":       label,
                    "type":        ntype,
                    "properties":  node.get("properties", {}),
                    "description": node.get("description", ""),
                }
            global_id = all_nodes[key]["id"]
            local_map[orig_id] = global_id

        for edge in g.get("edges", []):
            src = local_map.get(edge.get("source", ""), "")
            tgt = local_map.get(edge.get("target", ""), "")
            if src and tgt:
                all_edges.append({
                    "source":     src,
                    "target":     tgt,
                    "type":       edge.get("type", "RELAZIONATO_A"),
                    "label":      edge.get("label", ""),
                    "properties": edge.get("properties", {}),
                    "evidence":   edge.get("evidence", ""),
                })

    return {
        "nodes": list(all_nodes.values()),
        "edges": all_edges,
    }


# ─────────────────────────────────────────────────────────────────────────────
# 9. POST-PROCESSING: prune + dedup
# ─────────────────────────────────────────────────────────────────────────────

def prune_graph(graph: dict) -> dict:
    valid_ids = {n["id"] for n in graph["nodes"]}

    clean_edges = [
        e for e in graph["edges"]
        if e["source"] in valid_ids
        and e["target"] in valid_ids
        and e["source"] != e["target"]
    ]

    seen: set[tuple] = set()
    dedup_edges = []
    for e in clean_edges:
        key = (e["source"], e["target"], e["type"])
        if key not in seen:
            seen.add(key)
            dedup_edges.append(e)

    removed = len(graph["edges"]) - len(dedup_edges)
    if removed:
        print(f"   🧹  Prune: rimossi {removed} archi (orfani/duplicati/self-loop)")

    graph["edges"] = dedup_edges
    return graph


def llm_verify_relations(graph: dict, sample_size: int = 20) -> dict:
    if not graph["edges"]:
        return graph

    sample = graph["edges"][:sample_size]
    node_map = {n["id"]: n["label"] for n in graph["nodes"]}

    edges_text = "\n".join(
        f'{i}. {node_map.get(e["source"], "?")} --[{e["type"]}]--> '
        f'{node_map.get(e["target"], "?")}  | evidence: "{e.get("evidence","")}"'
        for i, e in enumerate(sample)
    )

    prompt = f"""Valuta le seguenti relazioni estratte da un documento.
Per ognuna indica se è valida in base all'evidence fornita.
Rispondi SOLO con JSON (array): [{{"idx":0,"valid":true,"reason":"..."}}]

Relazioni:
{edges_text}
"""
    try:
        raw = _call_llm([
                {"role": "system", "content": "Sei un revisore di knowledge graph. Rispondi solo con JSON array."},
                {"role": "user",   "content": prompt}
            ], 1000, temperature=0.0)
        verdicts = safe_parse_llm_json(raw) if raw else None
        if isinstance(verdicts, list):
            invalid_idx = {v["idx"] for v in verdicts if not v.get("valid", True)}
            before = len(graph["edges"])
            graph["edges"] = [
                e for i, e in enumerate(graph["edges"])
                if i >= sample_size or i not in invalid_idx
            ]
            removed = before - len(graph["edges"])
            if removed:
                print(f"   ✅  Verifica LLM: rimossi {removed} archi non validi")
    except Exception as e:
        print(f"   ⚠️  Verifica LLM fallita: {e}")

    return graph


GENERIC_TYPES = {
    "RELAZIONATO_A", "CONNESSO_A", "ASSOCIATO_A", "COLLEGATO_A", "MENZIONA",
    "HA_RELAZIONE", "E_CONNESSO", "APPARTIENE", "RIFERITO_A", "LEGATO_A",
    "RELATED_TO", "CONNECTED_TO", "ASSOCIATED_WITH", "LINKED_TO", "MENTIONS",
}


def enrich_relations(graph: dict, batch_size: int = 30) -> dict:
    node_map = {n["id"]: n["label"] for n in graph["nodes"]}

    to_enrich_idx = [
        i for i, e in enumerate(graph["edges"])
        if e.get("type", "").upper() in GENERIC_TYPES
        or not e.get("label", "").strip()
    ]

    if not to_enrich_idx:
        print("   ✅  Nessuna relazione generica trovata — arricchimento non necessario")
        return graph

    print(f"   🔧  {len(to_enrich_idx)} relazioni da arricchire…")
    enriched = 0

    for batch_start in range(0, len(to_enrich_idx), batch_size):
        batch_idx = to_enrich_idx[batch_start:batch_start + batch_size]
        batch_edges = []
        for i in batch_idx:
            e = graph["edges"][i]
            batch_edges.append({
                "_idx":     i,
                "source":   node_map.get(e["source"], e["source"]),
                "target":   node_map.get(e["target"], e["target"]),
                "type":     e.get("type", ""),
                "label":    e.get("label", ""),
                "properties": e.get("properties", {}),
                "evidence": e.get("evidence", ""),
            })

        payload = json.dumps({"edges": batch_edges}, ensure_ascii=False)
        try:
            raw = _call_llm([
                    {"role": "system", "content": ENRICH_PROMPT},
                    {"role": "user",   "content": payload},
                ], LLM_MAX_TOKENS)
            result = safe_parse_llm_json(raw) if raw else None

            if not result:
                continue

            improved = result.get("edges", result) if isinstance(result, dict) else result
            if not isinstance(improved, list):
                continue

            for item in improved:
                orig_idx = item.get("_idx")
                if orig_idx is None:
                    pos = improved.index(item)
                    if pos < len(batch_idx):
                        orig_idx = batch_idx[pos]
                    else:
                        continue

                e = graph["edges"][orig_idx]
                new_type  = item.get("type", e.get("type", "")).strip().upper().replace(" ", "_")
                new_label = item.get("label", e.get("label", "")).strip()
                new_props = item.get("properties", e.get("properties", {}))

                if new_type and new_type not in GENERIC_TYPES:
                    e["type"] = new_type
                    enriched += 1
                if new_label:
                    e["label"] = new_label
                if new_props:
                    e["properties"] = {**e.get("properties", {}), **new_props}

        except Exception as ex:
            print(f"   ⚠️  Errore arricchimento batch: {ex}")

    print(f"   ✅  Arricchimento completato: {enriched} relazioni migliorate")
    return graph


# ─────────────────────────────────────────────────────────────────────────────
# 10. CHECKPOINT
# ─────────────────────────────────────────────────────────────────────────────

def _checkpoint_path(input_path: Path) -> Path:
    return input_path.parent / (input_path.stem + ".checkpoint.json")


def _load_checkpoint(input_path: Path) -> list[dict] | None:
    cp = _checkpoint_path(input_path)
    if cp.exists():
        try:
            data = json.loads(cp.read_text(encoding="utf-8"))
            print(f"   ♻️  Checkpoint trovato → {cp.name} ({len(data)} grafi parziali, skip LLM)")
            return data
        except Exception:
            pass
    return None


def _save_checkpoint(input_path: Path, partial_graphs: list[dict]) -> None:
    cp = _checkpoint_path(input_path)
    cp.write_text(json.dumps(partial_graphs, ensure_ascii=False), encoding="utf-8")
    print(f"   💾  Checkpoint salvato → {cp.name}")


def _clear_checkpoint(input_path: Path) -> None:
    cp = _checkpoint_path(input_path)
    if cp.exists():
        cp.unlink()


# ─────────────────────────────────────────────────────────────────────────────
# 11. GENERAZIONE HTML
# ─────────────────────────────────────────────────────────────────────────────

# ─────────────────────────────────────────────────────────────────────────────
# 10. GENERAZIONE HTML
# ─────────────────────────────────────────────────────────────────────────────

NODE_COLORS = [
    "#4C8EDA", "#D75F5F", "#57C7E3", "#F16667", "#6DCE9E",
    "#FFC454", "#DA7194", "#845EC2", "#00C9A7", "#FF9671",
    "#F9F871", "#C34B96", "#2C73D2", "#0081CF", "#FF6F91",
]

def build_html(graph: dict, doc_name: str, output_path: Path) -> None:
    nodes = graph["nodes"]
    edges = graph["edges"]

    type_color: dict[str, str] = {}
    for node in nodes:
        t = node.get("type", "Nodo")
        if t not in type_color:
            idx = len(type_color) % len(NODE_COLORS)
            type_color[t] = NODE_COLORS[idx]

    js_nodes = []
    for node in nodes:
        color = type_color.get(node.get("type", "Nodo"), "#4C8EDA")
        js_nodes.append({
            "id":    node["id"],
            "label": node.get("label", ""),
            "type":  node.get("type", "Nodo"),
            "color": color,
            "desc":  node.get("description", ""),
            "props": node.get("properties", {}),
        })

    js_edges = []
    for edge in edges:
        js_edges.append({
            "source":   edge["source"],
            "target":   edge["target"],
            "type":     edge.get("type", ""),
            "label":    edge.get("label", ""),
            "props":    edge.get("properties", {}),
            "evidence": edge.get("evidence", ""),
        })

    legend_items = "".join(
        f'<div class="legend-item" data-type="{t}">'
        f'<span class="legend-dot" style="background:{c}"></span><span>{t}</span></div>'
        for t, c in sorted(type_color.items())
    )

    stats     = f"{len(nodes)} nodi &middot; {len(edges)} relazioni &middot; {len(type_color)} tipi"
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")

    cypher_lines = []
    for node in nodes:
        props = node.get("properties", {})
        ntype = node.get("type", "Nodo")
        label = node.get("label", "").replace("'", "\\'")
        prop_str = ", ".join(f"{k}: '{str(v).replace(chr(39), chr(92)+chr(39))}'" for k, v in props.items())
        if prop_str:
            cypher_lines.append(f"CREATE (:{ntype} {{id: '{node['id']}', name: '{label}', {prop_str}}});")
        else:
            cypher_lines.append(f"CREATE (:{ntype} {{id: '{node['id']}', name: '{label}'}});")
    for edge in edges:
        ev = edge.get("evidence", "").replace("'", "\\'")[:80]
        lbl = edge.get("label", "").replace("'", "\\'")
        eprops = edge.get("properties", {})
        prop_parts = [f"evidence:'{ev}'"]
        if lbl:
            prop_parts.append(f"label:'{lbl}'")
        for k, v in eprops.items():
            prop_parts.append(f"{k}:'{str(v).replace(chr(39), chr(92)+chr(39))}'")
        cypher_lines.append(
            f"MATCH (a {{id:'{edge['source']}'}}), (b {{id:'{edge['target']}'}}) "
            f"CREATE (a)-[:{edge['type']} {{{', '.join(prop_parts)}}}]->(b);"
        )
    cypher_str = "\n".join(cypher_lines)

    nodes_json  = json.dumps(js_nodes,  ensure_ascii=False)
    edges_json  = json.dumps(js_edges,  ensure_ascii=False)
    cypher_json = json.dumps(cypher_str, ensure_ascii=False)

    html = f"""<!DOCTYPE html>
<html lang="it">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Graph — {doc_name}</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
:root{{
  --bg:#1A1A2E;--panel:#16213E;--panel2:#0F3460;
  --accent:#E94560;--accent2:#F5A623;
  --text:#E0E0E0;--dim:#888;--border:#2A2A5A;
  --r:10px;--font:'Inter','Segoe UI',sans-serif;
}}
body{{background:var(--bg);color:var(--text);font-family:var(--font);
      height:100vh;display:flex;flex-direction:column;overflow:hidden}}
header{{display:flex;align-items:center;justify-content:space-between;
        padding:9px 18px;background:var(--panel);border-bottom:1px solid var(--border);flex-shrink:0}}
.hl{{display:flex;align-items:center;gap:10px}}
.logo{{font-size:17px;font-weight:700;color:#00C9A7;display:flex;align-items:center;gap:6px}}
.logo svg{{width:26px;height:26px}}
.dtitle{{font-size:12px;color:var(--dim);border-left:1px solid var(--border);padding-left:10px;max-width:300px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}
.badge{{font-size:11px;color:var(--dim);background:var(--panel2);padding:3px 10px;border-radius:20px}}
.hr{{display:flex;gap:7px;align-items:center}}
.btn{{cursor:pointer;padding:5px 13px;border:none;border-radius:6px;font-size:12px;font-weight:600;transition:.15s}}
.btn-a{{background:var(--accent);color:#fff}}.btn-a:hover{{background:#ff6a80}}
.btn-o{{background:transparent;border:1px solid var(--border);color:var(--text)}}.btn-o:hover{{background:var(--panel2)}}
.main{{display:flex;flex:1;overflow:hidden}}
.sb{{width:260px;min-width:200px;background:var(--panel);border-right:1px solid var(--border);
     display:flex;flex-direction:column;overflow:hidden;flex-shrink:0}}
.ss{{padding:12px 14px;border-bottom:1px solid var(--border)}}
.ss h3{{font-size:10px;text-transform:uppercase;letter-spacing:1px;color:var(--dim);margin-bottom:8px}}
.search{{width:100%;padding:6px 9px;background:var(--bg);border:1px solid var(--border);
         border-radius:6px;color:var(--text);font-size:12px;outline:none}}
.search:focus{{border-color:var(--accent2)}}
.cg{{display:grid;grid-template-columns:1fr 1fr 1fr;gap:5px}}
.cb{{padding:6px;font-size:11px;text-align:center;cursor:pointer;background:var(--bg);
     border:1px solid var(--border);border-radius:5px;color:var(--text);transition:.12s;user-select:none}}
.cb:hover{{background:var(--panel2);border-color:var(--accent2)}}
.cb-toggle.active{{background:var(--accent2);color:#111;border-color:var(--accent2)}}
.cb-toggle{{background:var(--bg);color:var(--dim)}}
.slider-row{{display:flex;align-items:center;gap:7px;margin-bottom:7px}}
.slider-row label{{font-size:11px;color:var(--dim);min-width:78px;flex-shrink:0}}
.slider-row span{{font-size:11px;color:var(--accent2);min-width:28px;text-align:right;flex-shrink:0}}
.sli{{flex:1;-webkit-appearance:none;height:4px;border-radius:2px;
      background:var(--border);outline:none;cursor:pointer}}
.sli::-webkit-slider-thumb{{-webkit-appearance:none;width:13px;height:13px;
      border-radius:50%;background:var(--accent2);cursor:pointer;transition:.1s}}
.sli::-webkit-slider-thumb:hover{{background:#ffd060}}
.sli::-moz-range-thumb{{width:13px;height:13px;border-radius:50%;
      background:var(--accent2);border:none;cursor:pointer}}
.ls{{overflow-y:auto;flex:1;padding:12px 14px}}
.legend-item{{display:flex;align-items:center;gap:7px;font-size:12px;margin-bottom:5px;
              cursor:pointer;padding:3px 5px;border-radius:4px;transition:.12s;user-select:none}}
.legend-item:hover,.legend-item.active{{background:var(--panel2)}}
.legend-dot{{width:11px;height:11px;border-radius:50%;flex-shrink:0}}
.ts{{padding:8px 14px;border-top:1px solid var(--border);font-size:10px;color:var(--dim)}}
#wrap{{flex:1;position:relative;overflow:hidden}}
canvas{{display:block;width:100%;height:100%;cursor:grab}}
canvas.dragging{{cursor:grabbing}}
#tip{{position:absolute;pointer-events:none;display:none;
      background:var(--panel);border:1px solid var(--border);border-radius:7px;
      padding:10px 12px;font-size:12px;max-width:260px;z-index:20;line-height:1.5}}
#tip strong{{color:var(--accent2);font-size:13px}}
#tip em{{color:var(--dim);font-size:11px}}
#tip .ev{{margin-top:5px;padding-top:5px;border-top:1px solid var(--border);
          font-size:11px;color:#aaa;font-style:italic}}
#info{{position:absolute;right:14px;bottom:14px;width:260px;
       background:var(--panel);border:1px solid var(--border);border-radius:var(--r);
       padding:13px;display:none;z-index:15;max-height:300px;overflow-y:auto}}
#info h4{{font-size:14px;font-weight:700;color:var(--accent2);margin-bottom:3px}}
#info .it{{font-size:11px;color:var(--dim);margin-bottom:8px}}
#info .ir{{font-size:12px;border-bottom:1px solid var(--border);padding:4px 0;display:flex;gap:8px}}
#info .ik{{color:var(--dim);min-width:75px}}
.xi{{float:right;cursor:pointer;color:var(--dim);font-size:16px;line-height:1}}
.mo{{display:none;position:fixed;inset:0;background:rgba(0,0,0,.75);z-index:100;
     align-items:center;justify-content:center}}
.mo.show{{display:flex}}
#ctx-menu{{
  position:fixed;display:none;z-index:200;
  background:var(--panel);border:1px solid var(--border);
  border-radius:8px;padding:5px 0;min-width:190px;
  box-shadow:0 8px 28px rgba(0,0,0,.6);font-size:13px;
}}
#ctx-menu .ctx-title{{
  padding:5px 14px 3px;font-size:10px;
  text-transform:uppercase;letter-spacing:1px;color:var(--dim);
  white-space:nowrap;overflow:hidden;text-overflow:ellipsis;max-width:180px;
}}
#ctx-menu .ctx-sep{{height:1px;background:var(--border);margin:4px 0}}
#ctx-menu .ctx-item{{
  display:flex;align-items:center;gap:9px;padding:7px 14px;
  cursor:pointer;color:var(--text);transition:.12s;user-select:none;
}}
#ctx-menu .ctx-item:hover{{background:var(--panel2);color:#fff}}
#ctx-menu .ctx-item.danger{{color:#ff8080}}
#ctx-menu .ctx-item.danger:hover{{background:#5a1010}}
.md{{background:var(--panel);border:1px solid var(--border);border-radius:var(--r);
     width:640px;max-width:94vw;max-height:78vh;display:flex;flex-direction:column}}
.mh{{display:flex;justify-content:space-between;align-items:center;
     padding:13px 16px;border-bottom:1px solid var(--border)}}
.mh h3{{font-size:13px;font-weight:700;color:var(--accent2)}}
.mb{{padding:14px;overflow-y:auto;flex:1}}
pre.cy{{background:var(--bg);border-radius:6px;padding:11px;font-size:11px;color:#00C9A7;
        overflow-x:auto;white-space:pre;line-height:1.6}}
.mf{{padding:9px 16px;border-top:1px solid var(--border);display:flex;justify-content:flex-end;gap:7px}}
footer{{font-size:10px;color:var(--dim);text-align:center;padding:4px;
        background:var(--panel);border-top:1px solid var(--border);flex-shrink:0}}
</style>
</head>
<body>

<header>
  <div class="hl">
    <div class="logo">
      <svg viewBox="0 0 60 60" xmlns="http://www.w3.org/2000/svg">
        <circle cx="14" cy="30" r="7" fill="#00C9A7"/>
        <circle cx="46" cy="13" r="7" fill="#F5A623"/>
        <circle cx="46" cy="47" r="7" fill="#E94560"/>
        <circle cx="30" cy="30" r="5" fill="#4C8EDA"/>
        <line x1="14" y1="30" x2="30" y2="30" stroke="#555" stroke-width="2"/>
        <line x1="30" y1="30" x2="46" y2="13" stroke="#555" stroke-width="2"/>
        <line x1="30" y1="30" x2="46" y2="47" stroke="#555" stroke-width="2"/>
        <line x1="46" y1="13" x2="46" y2="47" stroke="#555" stroke-width="2"/>
      </svg>
      Graph Explorer
    </div>
    <span class="dtitle">📄 {doc_name}</span>
  </div>
  <div class="hr">
    <span class="badge">{stats}</span>
    <button class="btn btn-o" onclick="openCypher()">Cypher</button>
    <button class="btn btn-a" onclick="resetView()">↺ Reset</button>
  </div>
</header>

<div class="main">
  <div class="sb">
    <div class="ss">
      <h3>🔍 Cerca nodo</h3>
      <input class="search" id="q" type="text" placeholder="Nome…" oninput="search(this.value)">
    </div>
    <div class="ss">
      <h3>⚙️ Controlli</h3>
      <div class="cg">
        <div class="cb" onclick="fitView()">Adatta</div>
        <div class="cb cb-toggle active" id="btn-phys" onclick="togglePhys()">Physics ON</div>
        <div class="cb" onclick="zoom(1.3)">Zoom +</div>
        <div class="cb" onclick="zoom(0.77)">Zoom −</div>
        <div class="cb" onclick="unpinAll()">Sblocca</div>
        <div class="cb" onclick="showAll()">Tutto</div>
        <div class="cb" onclick="savePNG()">PNG</div>
      </div>
    </div>
    <div class="ss">
      <h3>🔗 Fisica grafo</h3>
      <div class="slider-row">
        <label>Lung. archi</label>
        <input class="sli" id="sl-spring" type="range" min="40" max="500" value="150"
               oninput="setSpringL(+this.value)">
        <span id="lbl-spring">150</span>
      </div>
      <div class="slider-row">
        <label>Repulsione</label>
        <input class="sli" id="sl-rep" type="range" min="500" max="15000" step="500" value="5000"
               oninput="setRepulsion(+this.value)">
        <span id="lbl-rep">5000</span>
      </div>
      <div class="slider-row">
        <label>Gravità</label>
        <input class="sli" id="sl-grav" type="range" min="0" max="60" step="1" value="12"
               oninput="setGravity(+this.value)">
        <span id="lbl-grav">0.012</span>
      </div>
      <div class="slider-row">
        <label>Damping</label>
        <input class="sli" id="sl-damp" type="range" min="50" max="99" step="1" value="85"
               oninput="setDamp(+this.value)">
        <span id="lbl-damp">0.85</span>
      </div>
    </div>
    <h3 style="font-size:10px;text-transform:uppercase;letter-spacing:1px;color:var(--dim);padding:12px 14px 4px">🏷️ Tipi nodo</h3>
    <div class="ls">{legend_items}</div>
    <div class="ts">Generato: {timestamp}</div>
  </div>

  <div id="wrap">
    <canvas id="c"></canvas>
    <div id="tip"></div>
    <div id="info">
      <span class="xi" onclick="closeInfo()">×</span>
      <h4 id="i-label"></h4>
      <div class="it" id="i-type"></div>
      <div id="i-props"></div>
    </div>
  </div>
</div>

<footer>doc2graph · LLM Knowledge Graph Extractor · {timestamp}</footer>

<div id="ctx-menu">
  <div class="ctx-title" id="ctx-node-name">Nodo</div>
  <div class="ctx-sep"></div>
  <div class="ctx-item" onclick="ctxFilterNode()">&#128269; Solo questo nodo</div>
  <div class="ctx-item" onclick="ctxExpand(1)">&#127758; Espandi 1&#176; livello</div>
  <div class="ctx-item" onclick="ctxExpand(2)">&#127758; Espandi 2&#176; livello</div>
  <div class="ctx-sep"></div>
  <div class="ctx-item" onclick="ctxPinToggle()">&#128204; <span id="ctx-pin-lbl">Fissa posizione</span></div>
  <div class="ctx-item" onclick="ctxCopyLabel()">&#128203; Copia etichetta</div>
  <div class="ctx-sep"></div>
  <div class="ctx-item" onclick="showAll();closeCtx()">&#9851; Reset filtro</div>
  <div class="ctx-item danger" onclick="ctxHideNode()">&#128584; Nascondi nodo</div>
</div>

<div class="mo" id="cm">
  <div class="md">
    <div class="mh">
      <h3>🔵 Export Cypher (Neo4j)</h3>
      <button class="btn btn-o" style="padding:2px 8px;font-size:16px" onclick="closeCypher()">×</button>
    </div>
    <div class="mb"><pre class="cy" id="cy-text"></pre></div>
    <div class="mf">
      <button class="btn btn-o" onclick="copyCy()">📋 Copia</button>
      <button class="btn btn-a" onclick="dlCy()">⬇ .cypher</button>
    </div>
  </div>
</div>

<script>
const NODES_DATA = {nodes_json};
const EDGES_DATA = {edges_json};
const CYPHER     = {cypher_json};

const canvas = document.getElementById("c");
const ctx    = canvas.getContext("2d");
const DPR    = devicePixelRatio;

const nodes = NODES_DATA.map((d, i) => {{
  const angle = (2 * Math.PI * i) / NODES_DATA.length;
  const r     = Math.min(400, 80 + NODES_DATA.length * 7);
  return {{ ...d, x: Math.cos(angle)*r, y: Math.sin(angle)*r,
            vx:0, vy:0, hidden:false, highlighted:false }};
}});

const nodeMap = Object.fromEntries(nodes.map(n => [n.id, n]));
const edges   = EDGES_DATA.map(e => ({{
  ...e, source: nodeMap[e.source], target: nodeMap[e.target]
}})).filter(e => e.source && e.target);

let tx=0, ty=0, scale=1, physicsOn=true;

function resize() {{
  canvas.width  = canvas.offsetWidth  * DPR;
  canvas.height = canvas.offsetHeight * DPR;
  ctx.setTransform(1,0,0,1,0,0);
}}
window.addEventListener("resize", () => {{ resize(); fitView(); }});
resize();

function fitView() {{
  const vis = nodes.filter(n => !n.hidden);
  if (!vis.length) return;
  let minX=Infinity,maxX=-Infinity,minY=Infinity,maxY=-Infinity;
  vis.forEach(n=>{{ minX=Math.min(minX,n.x);maxX=Math.max(maxX,n.x);
                    minY=Math.min(minY,n.y);maxY=Math.max(maxY,n.y); }});
  const pw=canvas.width/DPR, ph=canvas.height/DPR;
  const gw=maxX-minX+120, gh=maxY-minY+120;
  scale=Math.min(pw/gw, ph/gh, 2)*0.9;
  tx=pw/2-(minX+maxX)/2*scale;
  ty=ph/2-(minY+maxY)/2*scale;
}}

// ── Controlli fisica da UI ──────────────────────────────────────────────────
function setSpringL(v){{
  SPRING_L=v;
  document.getElementById("lbl-spring").textContent=v;
}}
function setRepulsion(v){{
  REPULSION=v;
  document.getElementById("lbl-rep").textContent=v;
}}
function setGravity(v){{
  GRAVITY=v/1000;
  document.getElementById("lbl-grav").textContent=(v/1000).toFixed(3);
}}
function setDamp(v){{
  DAMP=v/100;
  document.getElementById("lbl-damp").textContent=(v/100).toFixed(2);
}}
let REPULSION = 5000;
let SPRING_K  = 0.03;
let SPRING_L  = 150;
let DAMP      = 0.85;
let GRAVITY   = 0.012;
// ── Colori per tipo di arco ──────────────────────────────────────────────────
const EDGE_PALETTE = [
  "#57C7E3","#6DCE9E","#FFC454","#DA7194","#C990C0",
  "#F79767","#4FC1E0","#A0D568","#FFCE54","#ED5565",
  "#AC92EC","#48CFAD","#FC6E51","#5D9CEC","#F6BB42",
];
const edgeColorMap = {{}};
let _edgePaletteIdx = 0;
function getEdgeColor(type) {{
  if (!edgeColorMap[type]) {{
    edgeColorMap[type] = EDGE_PALETTE[_edgePaletteIdx % EDGE_PALETTE.length];
    _edgePaletteIdx++;
  }}
  return edgeColorMap[type];
}}

// Precalcola bend index per archi paralleli (stesso source+target o inversi)
const edgeBendIndex = (() => {{
  const pairCount = {{}};
  const pairSeq   = {{}};
  edges.forEach(e => {{
    const key = [e.source.id, e.target.id].sort().join("||");
    pairCount[key] = (pairCount[key] || 0) + 1;
  }});
  edges.forEach(e => {{
    const key = [e.source.id, e.target.id].sort().join("||");
    pairSeq[key] = (pairSeq[key] || 0);
    e._bendIdx   = pairSeq[key];
    e._bendTotal = pairCount[key];
    pairSeq[key]++;
  }});
}})();

function drawEdge(e, sel) {{
  const [sx, sy] = worldToScreen(e.source.x, e.source.y);
  const [ex, ey] = worldToScreen(e.target.x, e.target.y);
  const dx = ex - sx, dy = ey - sy;
  const len = Math.sqrt(dx * dx + dy * dy);
  if (len < 2) return;

  const col  = sel ? "#FFA500" : getEdgeColor(e.type);
  const NODE_R = 22 * scale;
  const ux = dx / len, uy = dy / len;
  // punti di attacco al bordo dei cerchi
  const ax = sx + ux * NODE_R, ay = sy + uy * NODE_R;
  const bx = ex - ux * NODE_R, by = ey - uy * NODE_R;

  // curvatura: archi paralleli si divaricano, singoli curvano leggermente
  const total = e._bendTotal || 1;
  const idx   = e._bendIdx   || 0;
  // offset perpendicolare: distribuisce i paralleli da -max a +max
  const maxBend = 40 * scale;
  let bend;
  if (total === 1) {{
    bend = Math.min(28 * scale, len * 0.12); // curva leggera di default
  }} else {{
    bend = (idx - (total - 1) / 2) * (maxBend * 2 / Math.max(total - 1, 1));
  }}
  const px = -uy, py = ux; // perpendicolare
  const mx = (ax + bx) / 2, my = (ay + by) / 2;
  const cpx = mx + px * bend, cpy = my + py * bend;

  // ── linea curva ─────────────────────────────────────────────────────────
  ctx.beginPath();
  ctx.moveTo(ax * DPR, ay * DPR);
  ctx.quadraticCurveTo(cpx * DPR, cpy * DPR, bx * DPR, by * DPR);
  ctx.strokeStyle = col;
  ctx.lineWidth   = (sel ? 2.5 : 1.8) * DPR;
  ctx.globalAlpha = sel ? 1.0 : 0.75;
  ctx.stroke();
  ctx.globalAlpha = 1.0;

  // ── punta freccia alla fine della curva ──────────────────────────────────
  // tangente alla Bézier in t=1: direzione da cp a b
  const tx2 = bx - cpx, ty2 = by - cpy;
  const tlen = Math.sqrt(tx2 * tx2 + ty2 * ty2) || 1;
  const tux = tx2 / tlen, tuy = ty2 / tlen;
  const hw = (sel ? 7 : 5.5) * scale, hl = (sel ? 13 : 10) * scale;
  const perpx = -tuy, perpy = tux;
  ctx.beginPath();
  ctx.moveTo(bx * DPR, by * DPR);
  ctx.lineTo((bx - tux * hl + perpx * hw) * DPR, (by - tuy * hl + perpy * hw) * DPR);
  ctx.lineTo((bx - tux * hl - perpx * hw) * DPR, (by - tuy * hl - perpy * hw) * DPR);
  ctx.closePath();
  ctx.fillStyle   = col;
  ctx.globalAlpha = sel ? 1.0 : 0.85;
  ctx.fill();
  ctx.globalAlpha = 1.0;

  // ── label arco (pill con sfondo) ────────────────────────────────────────
  if (scale < 0.3) return; // troppo zoom-out per leggere

  // punto medio sulla curva (t=0.5 della quadratica)
  const lx = 0.25 * ax + 0.5 * cpx + 0.25 * bx;
  const ly = 0.25 * ay + 0.5 * cpy + 0.25 * by;

  const typeFS  = Math.max(9, Math.min(12, 11 * scale));
  const labelFS = Math.max(8, Math.min(10,  9 * scale));

  ctx.textAlign    = "center";
  ctx.textBaseline = "middle";

  // type: sfondo pill
  ctx.font = `600 ${{typeFS * DPR}}px Inter,sans-serif`;
  const tw  = ctx.measureText(e.type).width;
  const ph2 = typeFS * DPR * 1.4;
  const pw2 = tw + 10 * DPR;
  const pr  = ph2 / 2;
  // pill background
  ctx.beginPath();
  ctx.roundRect(
    (lx * DPR) - pw2 / 2, ly * DPR - ph2 / 2,
    pw2, ph2, pr
  );
  ctx.fillStyle   = sel ? "rgba(255,140,0,0.88)" : `${{col}}cc`;
  ctx.shadowColor = "rgba(0,0,0,0.5)";
  ctx.shadowBlur  = 4 * DPR;
  ctx.fill();
  ctx.shadowBlur = 0;
  ctx.fillStyle  = "#fff";
  ctx.fillText(e.type, lx * DPR, ly * DPR);

  // label leggibile sotto, solo se c'è spazio
  if (e.label && scale > 0.6) {{
    const labelY = ly + typeFS * scale * 1.6;
    ctx.font      = `italic ${{labelFS * DPR}}px Inter,sans-serif`;
    const lw2     = ctx.measureText(e.label).width + 8 * DPR;
    const lh2     = labelFS * DPR * 1.4;
    ctx.beginPath();
    ctx.roundRect(lx * DPR - lw2 / 2, labelY * DPR - lh2 / 2, lw2, lh2, lh2 / 2);
    ctx.fillStyle = "rgba(0,0,0,0.45)";
    ctx.fill();
    ctx.fillStyle = sel ? "#FFE0A0" : "#ddd";
    ctx.fillText(e.label, lx * DPR, labelY * DPR);
  }}
}}

function worldToScreen(x,y){{ return [x*scale+tx, y*scale+ty]; }}

function draw() {{
  ctx.clearRect(0, 0, canvas.width, canvas.height);

  // archi non selezionati prima, poi quelli selezionati sopra
  edges.forEach(e => {{
    if (e.source.hidden || e.target.hidden) return;
    const sel = !!(selectedNode && (e.source === selectedNode || e.target === selectedNode));
    if (!sel) drawEdge(e, false);
  }});
  edges.forEach(e => {{
    if (e.source.hidden || e.target.hidden) return;
    const sel = !!(selectedNode && (e.source === selectedNode || e.target === selectedNode));
    if (sel) drawEdge(e, true);
  }});
  const NODE_R=22;
  nodes.forEach(n=>{{
    if(n.hidden) return;
    const [sx,sy]=worldToScreen(n.x,n.y);
    const r=NODE_R*scale*DPR,sel=(n===selectedNode),hi=n.highlighted;
    if(sel){{
      ctx.beginPath();ctx.arc(sx*DPR,sy*DPR,r+6*DPR,0,Math.PI*2);
      ctx.fillStyle="rgba(255,165,0,0.25)";ctx.fill();
    }}
    ctx.beginPath();ctx.arc(sx*DPR,sy*DPR,r,0,Math.PI*2);
    ctx.fillStyle=hi?"#FFD700":sel?"#FF8C00":n.color;ctx.fill();
    ctx.strokeStyle=hi||sel?"#FFA500":"rgba(255,255,255,0.15)";
    ctx.lineWidth=(sel?3:1.5)*DPR;ctx.stroke();
    const fs2=Math.max(9,Math.min(13,13*scale));
    ctx.font=`bold ${{fs2*DPR}}px Inter,sans-serif`;
    ctx.fillStyle="#fff";ctx.textAlign="center";ctx.textBaseline="middle";
    let lbl=n.label;
    while(lbl.length>2&&ctx.measureText(lbl).width>r*1.7) lbl=lbl.slice(0,-1);
    if(lbl!==n.label) lbl=lbl.slice(0,-1)+"…";
    ctx.fillText(lbl,sx*DPR,sy*DPR);
    if(scale>0.5){{
      const fs3=Math.max(8,9*scale);
      ctx.font=`${{fs3*DPR}}px Inter,sans-serif`;
      ctx.fillStyle="rgba(255,255,255,0.55)";
      ctx.fillText(n.type,sx*DPR,(sy+NODE_R*scale+10)*DPR);
    }}
    // indicatore pin
    if(n.pinned){{
      ctx.beginPath();
      ctx.arc((sx+NODE_R*scale*0.7)*DPR,(sy-NODE_R*scale*0.7)*DPR,4*DPR,0,Math.PI*2);
      ctx.fillStyle="#F5A623";ctx.fill();
    }}
  }});
}}

// ── stato interazione ───────────────────────────────────────────────────────
let dragNode=null, dragOffX=0, dragOffY=0;
let panStart=null, panTx=0, panTy=0;
let selectedNode=null;
let mouseDownNode=null;   // nodo su cui è partito il mousedown
let mouseDownPos=null;    // posizione schermo del mousedown
let didDrag=false;        // true se il mouse si è spostato abbastanza → è un drag, non un click

const DRAG_THRESHOLD=4;  // pixel di movimento minimo per considerarlo drag

function loop() {{
  if(physicsOn) stepPhysics();
  draw();
  requestAnimationFrame(loop);
}}
loop();
setTimeout(fitView,400);

function screenToWorld(sx,sy){{ return [(sx-tx)/scale,(sy-ty)/scale]; }}

function nodeAt(sx,sy) {{
  const [wx,wy]=screenToWorld(sx,sy),R=22;
  return nodes.find(n=>!n.hidden&&Math.hypot(n.x-wx,n.y-wy)<R)||null;
}}

function edgeAt(sx,sy) {{
  const [wx,wy]=screenToWorld(sx,sy),THRESH=8;
  for(const e of edges){{
    if(e.source.hidden||e.target.hidden) continue;
    const dx=e.target.x-e.source.x,dy=e.target.y-e.source.y;
    const len=Math.sqrt(dx*dx+dy*dy);
    if(len<1) continue;
    const t=((wx-e.source.x)*dx+(wy-e.source.y)*dy)/(len*len);
    if(t<0||t>1) continue;
    const px=e.source.x+t*dx,py=e.source.y+t*dy;
    if(Math.hypot(wx-px,wy-py)<THRESH/scale) return e;
  }}
  return null;
}}

canvas.addEventListener("mousedown",e=>{{
  if(e.button!==0) return;
  const r=canvas.getBoundingClientRect();
  const sx=e.clientX-r.left, sy=e.clientY-r.top;
  const n=nodeAt(sx,sy);
  mouseDownPos={{x:sx,y:sy}};
  didDrag=false;
  if(n){{
    mouseDownNode=n;
    // prepara offset per il drag, ma non attiva ancora dragNode
    const [wx,wy]=screenToWorld(sx,sy);
    dragOffX=wx-n.x; dragOffY=wy-n.y;
  }} else {{
    mouseDownNode=null;
    panStart={{x:sx,y:sy}}; panTx=tx; panTy=ty;
  }}
}});

canvas.addEventListener("mousemove",e=>{{
  const r=canvas.getBoundingClientRect();
  const sx=e.clientX-r.left, sy=e.clientY-r.top;

  // Attiva drag solo dopo DRAG_THRESHOLD pixel di spostamento
  if(mouseDownNode && !dragNode){{
    const dist=Math.hypot(sx-mouseDownPos.x, sy-mouseDownPos.y);
    if(dist>DRAG_THRESHOLD){{
      dragNode=mouseDownNode;
      didDrag=true;
      canvas.classList.add("dragging");
    }}
  }}

  if(dragNode){{
    const [wx,wy]=screenToWorld(sx,sy);
    dragNode.x=wx-dragOffX;
    dragNode.y=wy-dragOffY;
    // pinned: azzera velocità così quando physics è ON non lo risucchia subito
    dragNode.vx=0; dragNode.vy=0;
    dragNode.pinned=true;
    document.getElementById("tip").style.display="none";
    return;
  }}

  if(panStart){{
    tx=panTx+(sx-panStart.x); ty=panTy+(sy-panStart.y);
    return;
  }}

  // tooltip hover
  const n=nodeAt(sx,sy);
  const tip=document.getElementById("tip");
  if(n){{
    let html=`<strong>${{n.label}}</strong> <em>[${{n.type}}]</em>`;
    if(n.pinned) html+=` <span style="color:#F5A623;font-size:10px">📌</span>`;
    if(n.desc) html+=`<br><span style="color:var(--dim)">${{n.desc}}</span>`;
    const pkeys=Object.keys(n.props||{{}});
    if(pkeys.length) html+="<br>"+pkeys.map(k=>`<b>${{k}}:</b> ${{n.props[k]}}`).join("<br>");
    tip.innerHTML=html; tip.style.display="block";
    tip.style.left=(sx+14)+"px"; tip.style.top=(sy-10)+"px";
  }} else {{
    const eg=edgeAt(sx,sy);
    if(eg){{
      const src=nodeMap[eg.source.id]||eg.source;
      const tgt=nodeMap[eg.target.id]||eg.target;
      let html=`<strong style="color:var(--accent2)">${{eg.type}}</strong>`;
      if(eg.label) html+=` <em style="color:#ccc;font-size:11px">— ${{eg.label}}</em>`;
      html+=`<br><span style="color:var(--dim);font-size:11px">${{src.label}} → ${{tgt.label}}</span>`;
      const epkeys=Object.keys(eg.props||{{}});
      if(epkeys.length){{
        html+=`<br>`+epkeys.map(k=>`<span style="color:var(--dim)">${{k}}:</span> <b>${{eg.props[k]}}</b>`).join(" &nbsp;");
      }}
      if(eg.evidence) html+=`<div class="ev">"${{eg.evidence.slice(0,120)}}${{eg.evidence.length>120?"…":""}}"</div>`;
      tip.innerHTML=html; tip.style.display="block";
      tip.style.left=(sx+14)+"px"; tip.style.top=(sy-10)+"px";
    }} else {{
      tip.style.display="none";
    }}
  }}
}});

canvas.addEventListener("mouseup",e=>{{
  canvas.classList.remove("dragging");
  dragNode=null;
  panStart=null;
  // Se NON era un drag → è un click: seleziona/deseleziona nodo
  if(!didDrag){{
    const r=canvas.getBoundingClientRect();
    const n=nodeAt(e.clientX-r.left, e.clientY-r.top);
    selectedNode=n||null;
    if(n) showInfo(n); else closeInfo();
  }}
  mouseDownNode=null;
}});

// Doppio click su nodo → toglie il pin (lo rilascia alla fisica)
canvas.addEventListener("dblclick",e=>{{
  const r=canvas.getBoundingClientRect();
  const n=nodeAt(e.clientX-r.left, e.clientY-r.top);
  if(n){{ n.pinned=false; n.vx=0; n.vy=0; }}
}});

canvas.addEventListener("wheel",e=>{{
  e.preventDefault();
  const r=canvas.getBoundingClientRect();
  const sx=e.clientX-r.left,sy=e.clientY-r.top;
  const factor=e.deltaY<0?1.15:1/1.15;
  const newScale=Math.max(0.1,Math.min(5,scale*factor));
  tx=sx-(sx-tx)/scale*newScale;ty=sy-(sy-ty)/scale*newScale;scale=newScale;
}},{{passive:false}});

let lastDist=0;
canvas.addEventListener("touchstart",e=>{{
  if(e.touches.length===2) lastDist=Math.hypot(
    e.touches[0].clientX-e.touches[1].clientX,
    e.touches[0].clientY-e.touches[1].clientY);
}},{{passive:true}});
canvas.addEventListener("touchmove",e=>{{
  if(e.touches.length===2){{
    const d=Math.hypot(e.touches[0].clientX-e.touches[1].clientX,
                       e.touches[0].clientY-e.touches[1].clientY);
    if(lastDist) scale=Math.max(0.1,Math.min(5,scale*d/lastDist));
    lastDist=d;
  }}
}},{{passive:true}});

function stepPhysics() {{
  const vis = nodes.filter(n => !n.hidden);
  if (vis.length < 2) return;
  for (let i = 0; i < vis.length; i++) {{
    for (let j = i + 1; j < vis.length; j++) {{
      const a = vis[i], b = vis[j];
      let dx = b.x - a.x;
      let dy = b.y - a.y;
      const d2 = dx * dx + dy * dy + 1;
      const f = REPULSION / d2;
      const d = Math.sqrt(d2);
      dx /= d; dy /= d;
      a.vx -= f * dx; a.vy -= f * dy;
      b.vx += f * dx; b.vy += f * dy;
    }}
  }}

  edges.forEach(e => {{
    if (e.source.hidden || e.target.hidden) return;
    const dx = e.target.x - e.source.x;
    const dy = e.target.y - e.source.y;
    const d = Math.sqrt(dx * dx + dy * dy) + 0.01;
    const f = SPRING_K * (d - SPRING_L);
    const fx = f * dx / d; const fy = f * dy / d;
    e.source.vx += fx; e.source.vy += fy;
    e.target.vx -= fx; e.target.vy -= fy;
  }});

  vis.forEach(n => {{ n.vx -= n.x * GRAVITY; n.vy -= n.y * GRAVITY; }});
  vis.forEach(n => {{
    if (n === dragNode || n.pinned) return;
    n.vx *= DAMP; n.vy *= DAMP;
    n.x += n.vx; n.y += n.vy;
  }});
}}

function showInfo(n) {{
  document.getElementById("i-label").textContent=n.label;
  document.getElementById("i-type").textContent="["+n.type+"]";
  const cc=edges.filter(e=>e.source===n||e.target===n).length;
  let rows=`<div class="ir"><span class="ik">Relazioni</span><span>${{cc}}</span></div>`;
  if(n.desc) rows+=`<div class="ir"><span class="ik">Desc.</span><span>${{n.desc}}</span></div>`;
  Object.entries(n.props||{{}}).forEach(([k,v])=>{{
    rows+=`<div class="ir"><span class="ik">${{k}}</span><span>${{v}}</span></div>`;
  }});
  document.getElementById("i-props").innerHTML=rows;
  document.getElementById("info").style.display="block";
}}
function closeInfo(){{ document.getElementById("info").style.display="none";selectedNode=null; }}

function zoom(f){{ scale=Math.max(0.1,Math.min(5,scale*f)); }}
function togglePhys(){{
  physicsOn=!physicsOn;
  const btn=document.getElementById("btn-phys");
  btn.textContent=physicsOn?"Physics ON":"Physics OFF";
  btn.classList.toggle("active",physicsOn);
  // quando si riattiva, azzera le velocità dei pinned così non esplodono
  if(physicsOn) nodes.forEach(n=>{{ if(n.pinned){{ n.vx=0;n.vy=0; }} }});
}}
function unpinAll(){{ nodes.forEach(n=>{{ n.pinned=false;n.vx=0;n.vy=0; }}); }}
function resetView(){{ fitView(); }}
function showAll(){{ nodes.forEach(n=>n.hidden=false);fitView(); }}

function search(q){{
  q=q.toLowerCase().trim();
  nodes.forEach(n=>{{ n.hidden=q?!n.label.toLowerCase().includes(q):false;n.highlighted=false; }});
  if(q) nodes.filter(n=>!n.hidden).forEach(n=>n.highlighted=true);
}}

document.querySelectorAll(".legend-item").forEach(el=>{{
  el.addEventListener("click",()=>{{
    el.classList.toggle("active");
    const active=[...document.querySelectorAll(".legend-item.active")].map(e=>e.dataset.type);
    if(!active.length){{ nodes.forEach(n=>n.hidden=false);return; }}
    nodes.forEach(n=>n.hidden=!active.includes(n.type));
    fitView();
  }});
}});

function savePNG(){{
  const a=document.createElement("a");
  a.href=canvas.toDataURL("image/png");a.download="graph.png";a.click();
}}
function openCypher(){{
  document.getElementById("cy-text").textContent=CYPHER;
  document.getElementById("cm").classList.add("show");
}}
function closeCypher(){{ document.getElementById("cm").classList.remove("show"); }}
function copyCy(){{ navigator.clipboard?.writeText(CYPHER); }}
function dlCy(){{
  const a=document.createElement("a");
  a.href="data:text/plain;charset=utf-8,"+encodeURIComponent(CYPHER);
  a.download="graph.cypher";a.click();
}}

// ── Context menu ────────────────────────────────────────────────────────────
let ctxNode=null;

function closeCtx(){{
  document.getElementById("ctx-menu").style.display="none";
  ctxNode=null;
}}

canvas.addEventListener("contextmenu",function(e){{
  e.preventDefault();
  var r=canvas.getBoundingClientRect();
  var n=nodeAt(e.clientX-r.left, e.clientY-r.top);
  if(!n){{ closeCtx(); return; }}
  ctxNode=n;
  selectedNode=n;
  showInfo(n);
  document.getElementById("ctx-node-name").textContent=n.label+" ["+n.type+"]";
  document.getElementById("ctx-pin-lbl").textContent=n.pinned?"Sblocca posizione":"Fissa posizione";
  var menu=document.getElementById("ctx-menu");
  menu.style.display="block";
  var mw=menu.offsetWidth, mh=menu.offsetHeight;
  var vw=window.innerWidth,  vh=window.innerHeight;
  menu.style.left=(e.clientX+mw>vw ? e.clientX-mw : e.clientX+4)+"px";
  menu.style.top =(e.clientY+mh>vh ? e.clientY-mh : e.clientY+4)+"px";
}});

document.addEventListener("click", closeCtx);
document.addEventListener("keydown",function(e){{ if(e.key==="Escape") closeCtx(); }});

function getNeighborIds(startNode, depth){{
  var visited=new Set([startNode.id]);
  var frontier=[startNode.id];
  for(var d=0;d<depth;d++){{
    var next=[];
    edges.forEach(function(e){{
      if(frontier.indexOf(e.source.id)!==-1 && !visited.has(e.target.id)){{
        visited.add(e.target.id); next.push(e.target.id);
      }}
      if(frontier.indexOf(e.target.id)!==-1 && !visited.has(e.source.id)){{
        visited.add(e.source.id); next.push(e.source.id);
      }}
    }});
    frontier=next;
    if(!frontier.length) break;
  }}
  return visited;
}}

function ctxFilterNode(){{
  if(!ctxNode) return;
  var keep=getNeighborIds(ctxNode,1);
  nodes.forEach(function(n){{ n.hidden=!keep.has(n.id); }});
  fitView(); closeCtx();
}}

function ctxExpand(depth){{
  if(!ctxNode) return;
  var keep=getNeighborIds(ctxNode,depth);
  nodes.forEach(function(n){{ n.hidden=!keep.has(n.id); }});
  fitView(); closeCtx();
}}

function ctxPinToggle(){{
  if(!ctxNode) return;
  ctxNode.pinned=!ctxNode.pinned;
  if(!ctxNode.pinned){{ ctxNode.vx=0; ctxNode.vy=0; }}
  closeCtx();
}}

function ctxCopyLabel(){{
  if(!ctxNode) return;
  if(navigator.clipboard) navigator.clipboard.writeText(ctxNode.label);
  closeCtx();
}}

function ctxHideNode(){{
  if(!ctxNode) return;
  ctxNode.hidden=true;
  if(selectedNode===ctxNode) closeInfo();
  closeCtx();
}}
</script>
</body>
</html>
"""
    output_path.write_text(html, encoding="utf-8")
    print(f"\n✅  HTML salvato → {output_path}")

# ─────────────────────────────────────────────────────────────────────────────
# 12. PIPELINE PRINCIPALE
# ─────────────────────────────────────────────────────────────────────────────

def process_single_file(input_path: Path) -> list[dict]:
    """
    Estrae e restituisce la lista di grafi parziali per un singolo file.
    Checkpoint granulare: salva dopo OGNI chunk, riprende dal chunk mancante.
    """
    print(f"\n  📂  {input_path.name}  ({input_path.stat().st_size // 1024} KB)")
    print(f"  {'─'*50}")

    # ── Checkpoint: carica grafi già estratti per questo file ────────────────
    cached = _load_checkpoint(input_path)  # lista di grafi parziali già fatti
    already_done = len(cached) if cached else 0

    text = extract_text(input_path)
    if not text.strip():
        print(f"  ⚠️  Nessun testo estratto da {input_path.name} — saltato")
        return []

    print(f"  ✅  {len(text):,} caratteri estratti")

    chunks = split_into_chunks(text, CHUNK_SIZE, CHUNK_OVERLAP)
    print(f"  ✅  {len(chunks)} chunk(s)  (già fatti: {already_done})")

    del text
    gc.collect()

    partial_graphs: list[dict] = cached if cached else []

    for i, chunk in enumerate(chunks, 1):
        if i <= already_done:
            continue  # chunk già in checkpoint, salta
        g = llm_extract_graph(chunk, i, len(chunks))
        partial_graphs.append(g)
        _save_checkpoint(input_path, partial_graphs)  # salva dopo ogni chunk
        print(f"   💾  Chunk {i}/{len(chunks)} salvato ({len(partial_graphs[-1].get('nodes',[]))} nodi, {len(partial_graphs[-1].get('edges',[]))} archi)")
        del chunk
        gc.collect()

    return partial_graphs


# ─────────────────────────────────────────────────────────────────────────────
# 13. PIPELINE MULTI-FILE
# ─────────────────────────────────────────────────────────────────────────────

def process_files(
    input_paths: list[Path],
    output_path: Path | None = None,
    verify: bool = False,
    enrich: bool = True,
) -> Path:
    """
    Processa uno o più file e genera un unico grafo HTML unificato.
    - Checkpoint per file: riprende se interrotto
    - Merge incrementale: bassa RAM (non accumula tutti i parziali)
    - gc.collect() tra un file e l'altro
    """
    n_files = len(input_paths)
    print(f"\n{'═'*60}")
    print(f"  📦  {n_files} file da processare")
    print(f"  🔌  LLM: {LLM_BASE_URL}  |  Modello: {LLM_MODEL or '(auto)'}")
    print(f"{'═'*60}")

    print("\n1/5  Estrazione testo e analisi LLM per ogni file…")

    # ── Merge incrementale: processa un file alla volta ──────────────────────
    accumulated: dict = {"nodes": [], "edges": []}

    for idx, path in enumerate(input_paths, 1):
        print(f"\n  [{idx}/{n_files}]", end="")
        partial = process_single_file(path)

        if partial:
            accumulated = merge_graphs([accumulated] + partial)
            accumulated = prune_graph(accumulated)
            print(f"   📊  Totale finora: {len(accumulated['nodes'])} nodi, "
                  f"{len(accumulated['edges'])} archi")

        del partial
        gc.collect()

    print(f"\n\n2/5  Merge completato: "
          f"{len(accumulated['nodes'])} nodi, {len(accumulated['edges'])} archi")

    # ── Arricchimento ────────────────────────────────────────────────────────
    print("\n3/5  Arricchimento relazioni…")
    if enrich:
        accumulated = enrich_relations(accumulated)
    else:
        print("   ⏭️  Saltato (--no-enrich attivo)")

    # ── Verifica opzionale ───────────────────────────────────────────────────
    if verify:
        print("\n4/5  Verifica relazioni via LLM…")
        accumulated = llm_verify_relations(accumulated)
    else:
        print("\n4/5  Verifica LLM saltata (usa --verify per attivarla)")

    # ── Output ───────────────────────────────────────────────────────────────
    print("\n5/5  Generazione HTML…")

    if n_files == 1:
        doc_name     = input_paths[0].name
        default_stem = input_paths[0].stem + "_graph"
        default_dir  = input_paths[0].parent
    else:
        names = ", ".join(p.name for p in input_paths[:3])
        if n_files > 3:
            names += f" +{n_files - 3} altri"
        doc_name     = names
        default_stem = "merged_graph"
        default_dir  = input_paths[0].parent

    if output_path is None:
        output_path = default_dir / (default_stem + ".html")

    build_html(accumulated, doc_name, output_path)

    json_out = output_path.with_suffix(".json")
    json_out.write_text(json.dumps(accumulated, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"   📊  JSON salvato → {json_out}")

    print(f"\n🎉  Completato!  {len(accumulated['nodes'])} nodi · {len(accumulated['edges'])} archi")
    print(f"    Apri nel browser:\n    {output_path.resolve()}\n")

    # Checkpoint puliti solo a successo completo
    for p in input_paths:
        _clear_checkpoint(p)

    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# 14. MODALITÀ --merge-jsons  (zero LLM)
# ─────────────────────────────────────────────────────────────────────────────

def merge_jsons(json_paths: list[Path], output_path: Path) -> Path:
    """
    Carica N file .json già estratti e li unisce in un HTML.
    Non chiama il LLM: utile per combinare batch precedenti.
    """
    print(f"\n🔀  Merge di {len(json_paths)} JSON (zero LLM)…")
    graphs = []
    for jp in json_paths:
        data = json.loads(jp.read_text(encoding="utf-8"))
        graphs.append(data)
        print(f"   ✅  {jp.name}: {len(data['nodes'])} nodi, {len(data['edges'])} archi")

    merged = merge_graphs(graphs)
    merged = prune_graph(merged)
    print(f"\n   📊  Dopo merge: {len(merged['nodes'])} nodi, {len(merged['edges'])} archi")

    doc_name = " + ".join(p.stem for p in json_paths[:3])
    if len(json_paths) > 3:
        doc_name += f" +{len(json_paths)-3}"

    build_html(merged, doc_name, output_path)
    json_out = output_path.with_suffix(".json")
    json_out.write_text(json.dumps(merged, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"\n✅  HTML → {output_path}")
    print(f"✅  JSON → {json_out}")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def main():
    global CHUNK_SIZE, CHUNK_OVERLAP

    parser = argparse.ArgumentParser(
        description=(
            "Estrae un grafo di conoscenza da uno o più documenti "
            "e genera un visualizzatore HTML interattivo stile Neo4j.\n\n"
            "Esempi:\n"
            "  python doc2graph.py relazione.pdf\n"
            "  python doc2graph.py file1.txt file2.docx file3.pdf\n"
            "  python doc2graph.py docs/*.pdf -o output/grafo.html\n"
            "  python doc2graph.py --merge-jsons batch1.json batch2.json -o finale.html"
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "inputs",
        nargs="*",
        help="Uno o più file (txt, md, pdf, docx, csv, json, epub…). "
             "Non richiesto in modalità --merge-jsons.",
    )
    parser.add_argument(
        "-o", "--output",
        help="File HTML di output",
    )
    parser.add_argument(
        "--chunk-size",
        type=int, default=CHUNK_SIZE,
        help=f"Dimensione chunk in caratteri (default: {CHUNK_SIZE})",
    )
    parser.add_argument(
        "--overlap",
        type=int, default=CHUNK_OVERLAP,
        help=f"Overlap tra chunk in caratteri (default: {CHUNK_OVERLAP})",
    )
    parser.add_argument(
        "--verify",
        action="store_true",
        help="Attiva verifica LLM delle relazioni estratte (costa token)",
    )
    parser.add_argument(
        "--no-enrich",
        action="store_true",
        help="Disattiva l'arricchimento automatico delle relazioni generiche",
    )
    parser.add_argument(
        "--merge-jsons",
        nargs="+",
        metavar="JSON",
        help="Modalità merge: unisce N file .json già estratti in un unico HTML (zero LLM). "
             "Esempio: --merge-jsons batch1.json batch2.json -o finale.html",
    )
    parser.add_argument(
        "--no-checkpoint",
        action="store_true",
        help="Disattiva il checkpointing per file (default: attivo)",
    )
    args = parser.parse_args()

    CHUNK_SIZE    = args.chunk_size
    CHUNK_OVERLAP = args.overlap

    # ── Modalità merge-jsons ─────────────────────────────────────────────────
    if args.merge_jsons:
        json_paths = []
        for raw in args.merge_jsons:
            p = Path(raw)
            if not p.exists():
                print(f"❌  File non trovato: {p}")
                sys.exit(1)
            json_paths.append(p)
        out = Path(args.output) if args.output else Path("merged_graph.html")
        merge_jsons(json_paths, out)
        return

    # ── Modalità normale ─────────────────────────────────────────────────────
    if not args.inputs:
        parser.print_help()
        sys.exit(1)

    input_paths: list[Path] = []
    for raw in args.inputs:
        p = Path(raw)
        if not p.exists():
            print(f"❌  File non trovato: {p}")
            sys.exit(1)
        if not p.is_file():
            print(f"❌  Non è un file: {p}")
            sys.exit(1)
        input_paths.append(p)

    if not input_paths:
        print("❌  Nessun file valido specificato.")
        sys.exit(1)

    # Se --no-checkpoint, cancella eventuali checkpoint esistenti
    if args.no_checkpoint:
        for p in input_paths:
            _clear_checkpoint(p)

    output_path = Path(args.output) if args.output else None

    process_files(
        input_paths,
        output_path,
        verify=args.verify,
        enrich=not args.no_enrich,
    )


if __name__ == "__main__":
    main()