"""
extractors.py — Estrazione testo da qualsiasi formato documento.

Supportati nativamente:
    .txt .md .log .rst .yaml .yml .toml .py .js .ts .java .cs .cpp .c .go
    .pdf  (pypdf + pdftotext fallback)
    .docx (python-docx + pandoc fallback)
    .csv / .tsv
    .json / .jsonl
    .epub .odt .rtf .doc .pptx  (pandoc)

Fallback universale: pandoc (se installato).
"""

from __future__ import annotations

import csv
import json
import subprocess
from pathlib import Path

# ── helpers ──────────────────────────────────────────────────────────────────

_TEXT_EXTS = {
    ".txt", ".md", ".log", ".rst", ".yaml", ".yml", ".toml",
    ".py", ".js", ".ts", ".java", ".cs", ".cpp", ".c", ".go",
    ".sh", ".bat", ".ini", ".cfg", ".xml", ".html", ".htm",
}


# ── public entry point ───────────────────────────────────────────────────────

def extract_text(path: Path) -> str:
    """Restituisce il testo estratto da *path*, qualunque sia il formato."""
    ext = path.suffix.lower()

    if ext in _TEXT_EXTS:
        return path.read_text(encoding="utf-8", errors="replace")

    dispatch = {
        ".pdf":  _extract_pdf,
        ".docx": _extract_docx,
        ".csv":  lambda p: _extract_csv(p, ".csv"),
        ".tsv":  lambda p: _extract_csv(p, ".tsv"),
        ".json": lambda p: _extract_json(p, ".json"),
        ".jsonl": lambda p: _extract_json(p, ".jsonl"),
    }

    if ext in dispatch:
        return dispatch[ext](path)

    if ext in {".epub", ".odt", ".rtf", ".doc", ".pptx"}:
        return _pandoc_to_text(path)

    print(f"⚠️  Formato '{ext}' non nativo — provo con pandoc…")
    return _pandoc_to_text(path)


# ── format-specific extractors ───────────────────────────────────────────────

def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
        reader = PdfReader(str(path))
        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append(f"--- Pagina {i + 1} ---\n{text}")
        full = "\n\n".join(pages)

        # Se il testo estratto è troppo scarso, prova pdftotext
        if len(full.strip()) < 200:
            result = subprocess.run(
                ["pdftotext", "-layout", str(path), "-"],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0 and result.stdout.strip():
                full = result.stdout

        return full

    except ImportError:
        print("  ⚠️  pypdf non installato — provo pdftotext…")
        return _pdftotext(path)
    except Exception as exc:
        return f"[Errore lettura PDF: {exc}]"


def _pdftotext(path: Path) -> str:
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", str(path), "-"],
            capture_output=True, text=True, timeout=60,
        )
        return result.stdout if result.returncode == 0 else f"[pdftotext error: {result.stderr[:200]}]"
    except FileNotFoundError:
        return "[pdftotext non trovato]"
    except Exception as exc:
        return f"[pdftotext error: {exc}]"


def _extract_docx(path: Path) -> str:
    try:
        import docx as _docx  # type: ignore
        doc = _docx.Document(str(path))
        parts: list[str] = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style = para.style.name if para.style else ""
            prefix = ""
            if "Heading" in style:
                level = style.replace("Heading ", "")
                prefix = "#" * int(level) + " " if level.isdigit() else "## "
            parts.append(prefix + para.text)

        for idx, table in enumerate(doc.tables):
            parts.append(f"\n[TABELLA {idx + 1}]")
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                sep = "| " + " | ".join(["---"] * len(table.columns)) + " |"
                parts += [rows[0], sep] + rows[1:]
            parts.append("")

        return "\n".join(parts)

    except ImportError:
        return _pandoc_to_text(path)
    except Exception:
        return _pandoc_to_text(path)


def _extract_csv(path: Path, ext: str) -> str:
    try:
        sep = "\t" if ext == ".tsv" else ","
        with open(path, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f, delimiter=sep)
            rows = list(reader)

        if not rows:
            return ""

        header = rows[0]
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * len(header)) + " |",
        ]
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")

        return f"[CSV — {len(rows) - 1} righe, {len(header)} colonne]\n" + "\n".join(lines)

    except Exception as exc:
        return f"[Errore lettura CSV: {exc}]"


def _extract_json(path: Path, ext: str) -> str:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        if ext == ".jsonl":
            items = []
            for line in text.splitlines():
                line = line.strip()
                if line:
                    try:
                        items.append(json.loads(line))
                    except json.JSONDecodeError:
                        pass
            return json.dumps(items[:100], indent=2, ensure_ascii=False)
        else:
            data = json.loads(text)
            return json.dumps(data, indent=2, ensure_ascii=False)
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")


def _pandoc_to_text(path: Path) -> str:
    try:
        result = subprocess.run(
            ["pandoc", str(path), "-t", "plain", "--wrap=none"],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0:
            return result.stdout
        return f"[Pandoc error: {result.stderr[:300]}]"
    except FileNotFoundError:
        return f"[pandoc non trovato — installa pandoc per supportare {path.suffix}]"
    except Exception as exc:
        return f"[Errore pandoc: {exc}]"
