from docx import Document

# Creiamo un nuovo documento Word
doc = Document()

# Titolo e introduzione
doc.add_heading("Modello Entità-Relazione per Sistema di Test Online", level=0)
doc.add_paragraph(
    "Questo documento descrive un modello Entità-Relazione complesso per un sistema di gestione "
    "di test online, comprendente studenti, insegnanti, corsi, test, domande, risposte, risultati "
    "e statistiche. Il modello è progettato per supportare analisi dettagliate e gestione avanzata."
)

# Sezione Tabella ER
doc.add_heading("Tabella Entità-Relazione", level=1)

# Definiamo i dati della tabella
entita_er = [
    ["Entità", "Attributi principali", "Relazioni", "Cardinalità"],
    ["Studente", "ID_Studente, Nome, Cognome, Email, Classe, Data_Nascita",
     "Partecipa a Test, Appartiene a Corso", "1:N Test, N:1 Corso"],
    ["Insegnante", "ID_Insegnante, Nome, Cognome, Email, Materia",
     "Crea Test, Gestisce Corso", "1:N Test, 1:N Corso"],
    ["Corso", "ID_Corso, Nome_Corso, Descrizione, Livello",
     "Contiene Studenti, Gestito da Insegnante", "1:N Studenti, 1:1 Insegnante"],
    ["Test", "ID_Test, Titolo, Data, Durata, Tipo (quiz/esame)",
     "Contiene Domande, Sostenuto da Studenti", "1:N Domande, N:M Studenti"],
    ["Domanda", "ID_Domanda, Testo, Tipo (scelta multipla, vero/falso, aperta), Punteggio",
     "Appartiene a Test, Ha Risposte", "1:N Risposte"],
    ["Risposta", "ID_Risposta, Testo, Corretta, Studente_Risposta",
     "Appartiene a Domanda, Fornita da Studente", "N:1 Domanda, N:1 Studente"],
    ["Risultato", "ID_Risultato, Punteggio, Data, Stato (passato/non passato)",
     "Collega Studente e Test", "1:1 Studente-Test"],
    ["Statistica", "ID_Statistica, Media_Punteggio, Percentuale_Corretta, Numero_Studenti",
     "Calcolata da Test", "1:1 Test"]
]

# Creiamo la tabella
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"

# Impostiamo l'intestazione
hdr_cells = table.rows[0].cells
for i, header in enumerate(entita_er[0]):
    hdr_cells[i].text = header

# Aggiungiamo le righe della tabella
for row_data in entita_er[1:]:
    row_cells = table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data

# Sezione Argomentazione
doc.add_heading("Argomentazione delle Entità e Relazioni", level=1)

argomentazione = [
    ("Studente", 
     "Memorizza tutte le informazioni necessarie per identificare e gestire gli studenti "
     "che partecipano ai test. Collegato ai corsi e ai test sostenuti."),
    ("Insegnante", 
     "Gestisce corsi e crea test, consentendo la pianificazione e valutazione delle attività didattiche."),
    ("Corso", 
     "Raggruppa studenti e collega insegnanti per materia, facilitando la gestione dei programmi e dei test."),
    ("Test", 
     "Oggetto centrale del sistema. Può essere di tipo quiz o esame, contiene domande e può essere sostenuto da più studenti."),
    ("Domanda", 
     "Dettaglia il contenuto dei test, includendo il testo della domanda, tipo e punteggio, permettendo valutazioni precise."),
    ("Risposta", 
     "Memorizza le risposte date dagli studenti a ciascuna domanda, permettendo analisi dettagliate dei risultati."),
    ("Risultato", 
     "Registra il punteggio ottenuto da uno studente in un test, con informazioni su data e stato del test."),
    ("Statistica", 
     "Aggrega dati dai test, fornendo medie, percentuali di risposte corrette e numero di partecipanti per analisi complessive.")
]

for entita, testo in argomentazione:
    doc.add_heading(entita, level=2)
    doc.add_paragraph(testo)

# Salviamo il documento
doc.save("Modello_ER_Sistema_Test.docx")

print("Documento creato: Modello_ER_Sistema_Test.docx")